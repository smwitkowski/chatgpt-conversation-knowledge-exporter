"""Tests for document parsing and normalization."""

from pathlib import Path
from tempfile import NamedTemporaryFile, TemporaryDirectory

import pytest

from ck_exporter.pipeline.io.documents import (
    generate_document_id,
    parse_docx_document,
    parse_markdown_document,
)
from ck_exporter.pipeline.io.load import load_conversations


def test_generate_document_id_stable():
    """Test that same content generates same ID."""
    path = Path("test_document.md")
    content1 = b"Test document content"
    content2 = b"Test document content"

    id1 = generate_document_id(path, content1)
    id2 = generate_document_id(path, content2)

    assert id1 == id2
    assert id1.startswith("doc__")
    assert "test_document" in id1


def test_generate_document_id_changes():
    """Test that different content generates different ID."""
    path = Path("test_document.md")
    content1 = b"Test document content"
    content2 = b"Different document content"

    id1 = generate_document_id(path, content1)
    id2 = generate_document_id(path, content2)

    assert id1 != id2
    assert id1.startswith("doc__")
    assert id2.startswith("doc__")


def test_generate_document_id_slug():
    """Test that filename is properly slugified."""
    path = Path("My Document 2025.md")
    content = b"content"

    doc_id = generate_document_id(path, content)

    assert "my-document-2025" in doc_id
    assert doc_id.startswith("doc__")


def test_parse_markdown_document_sections():
    """Test parsing Markdown document creates section messages."""
    content = """# Document Title

## Section One

This is section one content.

## Section Two

This is section two content.
"""

    with NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
        f.write(content)
        temp_path = Path(f.name)

    try:
        result = parse_markdown_document(temp_path)

        assert "conversation_id" in result
        assert result["conversation_id"].startswith("doc__")
        assert result["title"] == "Document Title"
        assert "mapping" in result
        assert "current_node" in result

        # Should have section messages
        mapping = result["mapping"]
        assert len(mapping) >= 2  # At least 2 sections

        # Check message IDs follow pattern sec:XXXX:slug
        for node_id in mapping.keys():
            assert node_id.startswith("sec:")
            assert ":" in node_id

        # Check first section has heading
        first_section = list(mapping.values())[0]
        assert "Section One" in first_section["message"]["content"]["parts"][0]
    finally:
        temp_path.unlink()


def test_parse_markdown_document_preface():
    """Test parsing Markdown with preface (content before first heading)."""
    content = """This is preface content before any heading.

## First Section

Section content here.
"""

    with NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
        f.write(content)
        temp_path = Path(f.name)

    try:
        result = parse_markdown_document(temp_path)

        mapping = result["mapping"]
        # Should have preface section
        preface_id = [nid for nid in mapping.keys() if "preface" in nid][0]
        preface_content = mapping[preface_id]["message"]["content"]["parts"][0]
        assert "preface content" in preface_content
    finally:
        temp_path.unlink()


def test_parse_markdown_document_empty():
    """Test parsing empty document creates fallback message."""
    content = ""

    with NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
        f.write(content)
        temp_path = Path(f.name)

    try:
        result = parse_markdown_document(temp_path)

        mapping = result["mapping"]
        # Should have a fallback document message
        assert len(mapping) > 0
        assert result["conversation_id"].startswith("doc__")
    finally:
        temp_path.unlink()


def test_parse_markdown_document_all_system_role():
    """Test that all document messages have system role."""
    content = """# Document

## Section

Content here.
"""

    with NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
        f.write(content)
        temp_path = Path(f.name)

    try:
        result = parse_markdown_document(temp_path)

        mapping = result["mapping"]
        for node in mapping.values():
            role = node["message"]["author"]["role"]
            assert role == "system"
    finally:
        temp_path.unlink()


@pytest.mark.importorskip("docx")
def test_parse_docx_document_sections():
    """Test parsing Word document creates section messages."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    with NamedTemporaryFile(suffix=".docx", delete=False) as f:
        temp_path = Path(f.name)

    try:
        # Create a simple Word document
        doc = Document()
        doc.add_heading("Document Title", 0)
        doc.add_heading("Section One", 1)
        doc.add_paragraph("This is section one content.")
        doc.add_heading("Section Two", 1)
        doc.add_paragraph("This is section two content.")
        doc.save(str(temp_path))

        result = parse_docx_document(temp_path)

        assert "conversation_id" in result
        assert result["conversation_id"].startswith("doc__")
        assert result["title"] == "Document Title"
        assert "mapping" in result
        assert "current_node" in result

        # Should have section messages
        mapping = result["mapping"]
        assert len(mapping) >= 2  # At least 2 sections

        # Check message IDs follow pattern sec:XXXX:slug
        for node_id in mapping.keys():
            assert node_id.startswith("sec:")
            assert ":" in node_id

        # Check first section has heading
        first_section = list(mapping.values())[0]
        assert "Section One" in first_section["message"]["content"]["parts"][0]
    finally:
        if temp_path.exists():
            temp_path.unlink()


def test_load_conversations_document_mode_md():
    """Test loading Markdown documents in document mode."""
    md_content = """# Document 1

## Section

Content here.
"""

    with TemporaryDirectory() as d:
        dir_path = Path(d)
        md_path = dir_path / "doc1.md"

        md_path.write_text(md_content, encoding="utf-8")

        result = load_conversations(dir_path, non_json_kind="document")

        assert len(result) == 1
        assert result[0]["conversation_id"].startswith("doc__")
        assert "mapping" in result[0]
        assert "current_node" in result[0]


def test_load_conversations_document_mode_skips_txt():
    """Test that .txt files are skipped in document mode."""
    md_content = """# Document

Content here.
"""
    txt_content = "Some text content"

    with TemporaryDirectory() as d:
        dir_path = Path(d)
        md_path = dir_path / "doc1.md"
        txt_path = dir_path / "doc2.txt"

        md_path.write_text(md_content, encoding="utf-8")
        txt_path.write_text(txt_content, encoding="utf-8")

        # Should only load .md, skip .txt
        result = load_conversations(dir_path, non_json_kind="document")

        assert len(result) == 1
        assert result[0]["conversation_id"].startswith("doc__")


def test_load_conversations_document_mode_mixed():
    """Test loading mixed .md and .docx files in document mode."""
    md_content = """# Document 1

Content here.
"""

    with TemporaryDirectory() as d:
        dir_path = Path(d)
        md_path = dir_path / "doc1.md"

        md_path.write_text(md_content, encoding="utf-8")

        # Try to create a docx if possible
        try:
            from docx import Document

            docx_path = dir_path / "doc2.docx"
            doc = Document()
            doc.add_heading("Document 2", 0)
            doc.add_paragraph("Content here.")
            doc.save(str(docx_path))

            result = load_conversations(dir_path, non_json_kind="document")
            assert len(result) == 2
            assert all(r["conversation_id"].startswith("doc__") for r in result)
        except ImportError:
            # If docx not available, just test md
            result = load_conversations(dir_path, non_json_kind="document")
            assert len(result) == 1


def test_load_conversations_meeting_mode_default():
    """Test that default mode (meeting) still works for .md files."""
    md_content = """# Meeting Notes

## Summary

Meeting summary here.
"""

    with TemporaryDirectory() as d:
        dir_path = Path(d)
        md_path = dir_path / "meeting1.md"

        md_path.write_text(md_content, encoding="utf-8")

        result = load_conversations(dir_path)  # Default is meeting mode

        assert len(result) == 1
        assert result[0]["conversation_id"].startswith("meeting__")

