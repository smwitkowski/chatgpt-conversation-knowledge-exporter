"""Parsing utilities for generic documents (Markdown and Word documents).

Supports:
- Markdown documents (.md): Split by headings into sections
- Word documents (.docx): Split by heading styles into sections
"""

import hashlib
import re
from pathlib import Path
from typing import Dict, List, Optional

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore


def _slugify(stem: str) -> str:
    """Convert a string to a URL-safe slug."""
    slug = re.sub(r"[^\w\-]", "-", stem.lower())
    slug = re.sub(r"-+", "-", slug).strip("-")
    return slug or "document"


def generate_document_id(path: Path, content: bytes) -> str:
    """
    Generate a stable document ID from filename and content hash.

    Format: `doc__<slug>__<hash8>`

    Args:
        path: File path
        content: File content bytes

    Returns:
        Stable document ID
    """
    # Create slug from filename stem
    slug = _slugify(path.stem)

    # Compute content hash (first 8 chars)
    content_hash = hashlib.sha256(content).hexdigest()[:8]

    return f"doc__{slug}__{content_hash}"


def _split_markdown_sections(text: str) -> List[tuple[str, str]]:
    """
    Split markdown text into sections by headings.

    Returns:
        List of (heading, body) tuples where heading may be "" for preface.
    """
    lines = text.split("\n")
    sections: List[dict] = []
    current = {"heading": "", "content": []}

    for line in lines:
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            sections.append(current)
            current = {"heading": m.group(2).strip(), "content": []}
        else:
            current["content"].append(line)

    sections.append(current)

    out: List[tuple[str, str]] = []
    for s in sections:
        heading = s["heading"].strip()
        body = "\n".join(s["content"]).strip()
        if not heading and not body:
            continue
        out.append((heading, body))
    return out


def parse_markdown_document(path: Path) -> Dict:
    """
    Parse a Markdown document file into a synthetic conversation dict.

    Splits the document by headings:
    - Each heading section becomes a "system" message with ID like "sec:0001:heading-slug"
    - Preface (content before first heading) becomes "sec:0001:preface"

    Args:
        path: Path to Markdown file

    Returns:
        Synthetic conversation dict with mapping/current_node structure
    """
    content_bytes = path.read_bytes()
    content = content_bytes.decode("utf-8", errors="replace")
    doc_id = generate_document_id(path, content_bytes)

    # Extract title from first heading or filename
    title = path.stem
    for line in content.split("\n")[:20]:  # Check first 20 lines for title
        if line.startswith("#"):
            title = re.sub(r"^#+\s*", "", line).strip() or title
            break

    sections = _split_markdown_sections(content)

    mapping = {}
    prev: Optional[str] = None
    for idx, (heading, body) in enumerate(sections, start=1):
        # Stable message id per section
        sec_slug = _slugify(heading) if heading else "preface"
        message_id = f"sec:{idx:04d}:{sec_slug}"

        full_text = heading
        if body:
            full_text = f"{heading}\n\n{body}" if heading else body
        if not full_text.strip():
            continue

        mapping[message_id] = {
            "id": message_id,
            "parent": prev,
            "message": {
                "id": message_id,
                "author": {"role": "system", "name": None, "metadata": {}},
                "create_time": None,
                "update_time": None,
                "content": {"content_type": "text", "parts": [full_text]},
                "status": "finished_successfully",
                "end_turn": True,
                "weight": 1,
                "metadata": {},
                "recipient": "all",
                "channel": None,
            },
        }
        prev = message_id

    # If no sections were created, create a single message with all content
    if not mapping:
        message_id = "sec:0001:document"
        mapping[message_id] = {
            "id": message_id,
            "parent": None,
            "message": {
                "id": message_id,
                "author": {"role": "system", "name": None, "metadata": {}},
                "create_time": None,
                "update_time": None,
                "content": {"content_type": "text", "parts": [content]},
                "status": "finished_successfully",
                "end_turn": True,
                "weight": 1,
                "metadata": {},
                "recipient": "all",
                "channel": None,
            },
        }
        prev = message_id

    current_node = prev if prev else None

    return {
        "conversation_id": doc_id,
        "title": title or "Untitled Document",
        "mapping": mapping,
        "current_node": current_node,
    }


def parse_docx_document(path: Path) -> Dict:
    """
    Parse a Word document (.docx) file into a synthetic conversation dict.

    Splits the document by headings (similar to parse_markdown_document):
    - Each heading section becomes a "system" message with ID like "sec:0001:heading-slug"
    - Preface (content before first heading) becomes "sec:0001:preface"

    Args:
        path: Path to .docx file

    Returns:
        Synthetic conversation dict with mapping/current_node structure

    Raises:
        ImportError: If python-docx is not installed
    """
    if Document is None:
        raise ImportError(
            "python-docx is required to parse .docx files. "
            "Install it with: pip install python-docx"
        )

    content_bytes = path.read_bytes()
    doc = Document(path)
    doc_id = generate_document_id(path, content_bytes)

    # Extract title from first heading or filename
    title = path.stem
    for para in doc.paragraphs[:20]:  # Check first 20 paragraphs for title
        style = para.style.name if para.style else "normal"
        text = para.text.strip()
        if text and (style.startswith("Heading") or style == "Title"):
            title = text
            break

    # Build sections by headings (simple, deterministic)
    sections: List[dict] = []
    current = {"heading": "", "content": []}

    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style else "normal"

        if not text:
            if current["content"]:
                current["content"].append("")
            continue

        if style.startswith("Heading") or style == "Title":
            sections.append(current)
            current = {"heading": text, "content": []}
        else:
            current["content"].append(text)

    sections.append(current)

    mapping = {}
    prev: Optional[str] = None
    sec_idx = 0
    for s in sections:
        heading = (s["heading"] or "").strip()
        body = "\n".join(s["content"]).strip()
        if not heading and not body:
            continue

        sec_idx += 1
        sec_slug = _slugify(heading) if heading else "preface"
        message_id = f"sec:{sec_idx:04d}:{sec_slug}"

        full_text = heading
        if body:
            full_text = f"{heading}\n\n{body}" if heading else body

        mapping[message_id] = {
            "id": message_id,
            "parent": prev,
            "message": {
                "id": message_id,
                "author": {"role": "system", "name": None, "metadata": {}},
                "create_time": None,
                "update_time": None,
                "content": {"content_type": "text", "parts": [full_text]},
                "status": "finished_successfully",
                "end_turn": True,
                "weight": 1,
                "metadata": {},
                "recipient": "all",
                "channel": None,
            },
        }
        prev = message_id

    current_node = prev if prev else None

    return {
        "conversation_id": doc_id,
        "title": title or "Untitled Document",
        "mapping": mapping,
        "current_node": current_node,
    }

