"""Parsing utilities for meeting notes and transcripts.

Supports:
- Markdown meeting notes (Google Meet style with notes + transcript)
- Plain text transcripts (Teams style with timestamped lines)
- Word documents (.docx) with meeting notes (Google Meet style with notes + transcript)
"""

import hashlib
import re
from pathlib import Path
from typing import Dict, List, Optional
from urllib.parse import quote

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore


def normalize_timestamp(raw: str) -> str:
    """
    Normalize a timestamp string to HH:MM:SS format.

    Handles:
    - M:SS format (e.g., "1:08" -> "00:01:08")
    - H:MM:SS format (e.g., "1:02:15" -> "01:02:15")
    - HH:MM:SS format (already normalized)
    - Minutes >= 60 are carried into hours (e.g., "62:15" -> "01:02:15")

    Args:
        raw: Raw timestamp string

    Returns:
        Normalized timestamp in HH:MM:SS format
    """
    if not raw:
        return "00:00:00"

    # Remove any anchor links or extra formatting
    raw = raw.strip()
    # Remove anchor link syntax like {#00:00:00}
    raw = re.sub(r'\s*\{#[^}]+\}', '', raw)

    # Try to parse as H:MM:SS or HH:MM:SS
    hmmss_match = re.match(r'^(\d{1,2}):(\d{2}):(\d{2})$', raw)
    if hmmss_match:
        hours = int(hmmss_match.group(1))
        minutes = int(hmmss_match.group(2))
        seconds = int(hmmss_match.group(3))
    else:
        # Try to parse as M:SS or MM:SS
        mss_match = re.match(r'^(\d{1,3}):(\d{2})$', raw)
        if mss_match:
            minutes = int(mss_match.group(1))
            seconds = int(mss_match.group(2))
            hours = 0
        else:
            # Invalid format, return default
            return "00:00:00"

    # Carry minutes >= 60 into hours
    hours += minutes // 60
    minutes = minutes % 60

    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"


def generate_document_id(path: Path, content: bytes) -> str:
    """
    Generate a stable document ID from filename and content hash.

    Format: `meeting__<slug>__<hash8>`

    Args:
        path: File path
        content: File content bytes

    Returns:
        Stable document ID
    """
    # Create slug from filename stem
    stem = path.stem
    # Replace spaces and special chars with hyphens, lowercase
    slug = re.sub(r'[^\w\-]', '-', stem.lower())
    slug = re.sub(r'-+', '-', slug).strip('-')

    # Compute content hash (first 8 chars)
    content_hash = hashlib.sha256(content).hexdigest()[:8]

    return f"meeting__{slug}__{content_hash}"


def _slugify_heading(heading: str) -> str:
    """Convert a heading to a URL-safe slug."""
    # Remove markdown heading markers
    heading = re.sub(r'^#+\s*', '', heading)
    # Replace spaces and special chars with hyphens
    slug = re.sub(r'[^\w\-]', '-', heading.lower())
    slug = re.sub(r'-+', '-', slug).strip('-')
    return slug or "section"


def _is_timestamp_heading(heading: str) -> Optional[str]:
    """
    Check if a heading contains a timestamp pattern.

    Returns the timestamp string if found, None otherwise.
    """
    # Look for patterns like "### 00:03:03" or "### 00:03:03 {#00:03:03}"
    match = re.search(r'(\d{1,2}:\d{2}(?::\d{2})?)', heading)
    if match:
        return match.group(1)
    return None


def parse_markdown_meeting(path: Path) -> Dict:
    """
    Parse a Markdown meeting notes file into a synthetic conversation dict.

    Splits the document by headings:
    - Notes sections become "system" messages with IDs like "notes:summary"
    - Transcript timestamp sections become messages with normalized timestamp IDs

    Args:
        path: Path to Markdown file

    Returns:
        Synthetic conversation dict with mapping/current_node structure
    """
    content_bytes = path.read_bytes()
    content = content_bytes.decode('utf-8')
    doc_id = generate_document_id(path, content_bytes)

    # Extract title from first heading or filename
    title = path.stem
    lines = content.split('\n')
    for line in lines[:20]:  # Check first 20 lines for title
        if line.startswith('#'):
            title = re.sub(r'^#+\s*', '', line).strip()
            break

    # Split content by headings
    sections = []
    current_section = {"heading": None, "level": 0, "content": []}

    for line in lines:
        # Check if this is a heading
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            # Save previous section if it has content
            if current_section["heading"] is not None or current_section["content"]:
                sections.append(current_section)

            # Start new section
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2)
            current_section = {
                "heading": heading_text,
                "level": level,
                "content": []
            }
        else:
            current_section["content"].append(line)

    # Add final section
    if current_section["heading"] is not None or current_section["content"]:
        sections.append(current_section)

    # Build mapping structure
    mapping = {}
    previous_node_id = None

    for section in sections:
        if not section["heading"] and not section["content"]:
            continue

        heading = section["heading"] or ""
        section_text = '\n'.join(section["content"]).strip()

        # Check if this is a transcript timestamp section
        timestamp = _is_timestamp_heading(heading)
        if timestamp:
            # This is a transcript message
            message_id = normalize_timestamp(timestamp)
            role = "user"  # Transcript messages are user-like
        else:
            # This is a notes section
            slug = _slugify_heading(heading) if heading else "preface"
            message_id = f"notes:{slug}"
            role = "system"

            # Special case: prepend action items hint for checklist sections
            heading_lower = heading.lower()
            if ("next steps" in heading_lower or "action" in heading_lower or
                "todo" in heading_lower or "tasks" in heading_lower):
                section_text = f"Action items (treat as commitments/tasks):\n\n{section_text}"

        # Combine heading and content
        if heading:
            full_text = f"{heading}\n\n{section_text}" if section_text else heading
        else:
            full_text = section_text

        if not full_text.strip():
            continue

        # Create message node
        node_id = message_id
        mapping[node_id] = {
            "id": node_id,
            "parent": previous_node_id,
            "message": {
                "id": message_id,
                "author": {
                    "role": role,
                    "name": None,
                    "metadata": {},
                },
                "create_time": None,  # No absolute time for v1
                "update_time": None,
                "content": {
                    "content_type": "text",
                    "parts": [full_text],
                },
                "status": "finished_successfully",
                "end_turn": True,
                "weight": 1,
                "metadata": {},
                "recipient": "all",
                "channel": None,
            },
        }

        previous_node_id = node_id

    # Current node is the last message
    current_node = previous_node_id if previous_node_id else None

    return {
        "conversation_id": doc_id,
        "title": title or "Untitled Meeting",
        "mapping": mapping,
        "current_node": current_node,
    }


def parse_text_transcript(path: Path) -> Dict:
    """
    Parse a plain text transcript file into a synthetic conversation dict.

    Expects lines in format: `TIME : NAME : TEXT`

    Args:
        path: Path to text file

    Returns:
        Synthetic conversation dict with mapping/current_node structure
    """
    content_bytes = path.read_bytes()
    content = content_bytes.decode('utf-8')
    doc_id = generate_document_id(path, content_bytes)

    # Extract title from filename
    title = path.stem

    lines = content.split('\n')
    mapping = {}
    previous_node_id = None

    # Pattern to match: TIME : NAME : TEXT
    # Examples: "1:08 : Tanya Gastelum : Good, good afternoon."
    pattern = re.compile(r'^(\d{1,3}:\d{2}(?::\d{2})?)\s*:\s*([^:]+?)\s*:\s*(.+)$')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        match = pattern.match(line)
        if match:
            time_str = match.group(1)
            name = match.group(2).strip()
            text = match.group(3).strip()

            # Normalize timestamp
            message_id = normalize_timestamp(time_str)

            # Format as markdown-style speaker line
            formatted_text = f"**{name}:** {text}"

            # Create message node
            node_id = message_id
            mapping[node_id] = {
                "id": node_id,
                "parent": previous_node_id,
                "message": {
                    "id": message_id,
                    "author": {
                        "role": "user",
                        "name": None,
                        "metadata": {},
                    },
                    "create_time": None,  # No absolute time for v1
                    "update_time": None,
                    "content": {
                        "content_type": "text",
                        "parts": [formatted_text],
                    },
                    "status": "finished_successfully",
                    "end_turn": True,
                    "weight": 1,
                    "metadata": {},
                    "recipient": "all",
                    "channel": None,
                },
            }

            previous_node_id = node_id
        else:
            # Non-matching line: attach to previous message as continuation
            if previous_node_id and previous_node_id in mapping:
                existing_parts = mapping[previous_node_id]["message"]["content"]["parts"]
                if existing_parts:
                    existing_parts[-1] += f"\n{line}"
                else:
                    existing_parts.append(line)

    # If no messages were created, create a single message with all content
    if not mapping:
        node_id = "notes:transcript"
        mapping[node_id] = {
            "id": node_id,
            "parent": None,
            "message": {
                "id": node_id,
                "author": {
                    "role": "system",
                    "name": None,
                    "metadata": {},
                },
                "create_time": None,
                "update_time": None,
                "content": {
                    "content_type": "text",
                    "parts": [content],
                },
                "status": "finished_successfully",
                "end_turn": True,
                "weight": 1,
                "metadata": {},
                "recipient": "all",
                "channel": None,
            },
        }
        previous_node_id = node_id

    current_node = previous_node_id if previous_node_id else None

    return {
        "conversation_id": doc_id,
        "title": title or "Untitled Transcript",
        "mapping": mapping,
        "current_node": current_node,
    }


def parse_docx_meeting(path: Path) -> Dict:
    """
    Parse a Word document (.docx) meeting notes file into a synthetic conversation dict.

    Splits the document by headings (similar to parse_markdown_meeting):
    - Notes sections become "system" messages with IDs like "notes:summary"
    - Transcript timestamp sections become messages with normalized timestamp IDs

    Args:
        path: Path to .docx file

    Returns:
        Synthetic conversation dict with mapping/current_node structure

    Raises:
        ImportError: If python-docx is not installed
    """
    if Document is None:
        raise ImportError(
            "python-docx is required to parse .docx files. "
            "Install it with: pip install python-docx"
        )

    content_bytes = path.read_bytes()
    doc = Document(path)
    doc_id = generate_document_id(path, content_bytes)

    # Extract title from first heading or filename
    title = path.stem
    for para in doc.paragraphs[:20]:  # Check first 20 paragraphs for title
        style = para.style.name if para.style else 'normal'
        text = para.text.strip()
        if style.startswith('Heading') or style == 'Title':
            title = text
            break

    # Split content by headings
    sections = []
    current_section = {"heading": None, "level": 0, "content": []}

    for para in doc.paragraphs:
        style = para.style.name if para.style else 'normal'
        text = para.text.strip()

        # Skip empty paragraphs but preserve them in content for spacing
        if not text:
            if current_section["content"]:
                current_section["content"].append("")
            continue

        # Check if it's a heading
        if style.startswith('Heading'):
            # Save previous section if it has content
            if current_section["heading"] is not None or current_section["content"]:
                sections.append(current_section)

            # Extract heading level
            level = int(style.split()[-1]) if style.split()[-1].isdigit() else 1
            current_section = {
                "heading": text,
                "level": level,
                "content": []
            }
        elif style == 'Title':
            # Title is like a top-level heading
            if current_section["heading"] is not None or current_section["content"]:
                sections.append(current_section)
            current_section = {
                "heading": text,
                "level": 1,
                "content": []
            }
        else:
            current_section["content"].append(text)

    # Add final section
    if current_section["heading"] is not None or current_section["content"]:
        sections.append(current_section)

    # Build mapping structure (same logic as parse_markdown_meeting)
    mapping = {}
    previous_node_id = None

    for section in sections:
        if not section["heading"] and not section["content"]:
            continue

        heading = section["heading"] or ""
        section_text = '\n'.join(section["content"]).strip()

        # Check if this is a transcript timestamp section
        timestamp = _is_timestamp_heading(heading)
        if timestamp:
            # This is a transcript message
            message_id = normalize_timestamp(timestamp)
            role = "user"  # Transcript messages are user-like
        else:
            # This is a notes section
            slug = _slugify_heading(heading) if heading else "preface"
            message_id = f"notes:{slug}"
            role = "system"

            # Special case: prepend action items hint for checklist sections
            heading_lower = heading.lower()
            if ("next steps" in heading_lower or "action" in heading_lower or
                "todo" in heading_lower or "tasks" in heading_lower):
                section_text = f"Action items (treat as commitments/tasks):\n\n{section_text}"

        # Combine heading and content
        if heading:
            full_text = f"{heading}\n\n{section_text}" if section_text else heading
        else:
            full_text = section_text

        if not full_text.strip():
            continue

        # Create message node
        node_id = message_id
        mapping[node_id] = {
            "id": node_id,
            "parent": previous_node_id,
            "message": {
                "id": message_id,
                "author": {
                    "role": role,
                    "name": None,
                    "metadata": {},
                },
                "create_time": None,  # No absolute time for v1
                "update_time": None,
                "content": {
                    "content_type": "text",
                    "parts": [full_text],
                },
                "status": "finished_successfully",
                "end_turn": True,
                "weight": 1,
                "metadata": {},
                "recipient": "all",
                "channel": None,
            },
        }

        previous_node_id = node_id

    # Current node is the last message
    current_node = previous_node_id if previous_node_id else None

    return {
        "conversation_id": doc_id,
        "title": title or "Untitled Meeting",
        "mapping": mapping,
        "current_node": current_node,
    }


def is_meeting_artifact(path: Path) -> bool:
    """
    Check if a file path is a meeting artifact (Markdown, text, or Word document).

    Args:
        path: File path to check

    Returns:
        True if file is .md, .txt, or .docx, False otherwise
    """
    return path.suffix.lower() in ('.md', '.txt', '.docx')

